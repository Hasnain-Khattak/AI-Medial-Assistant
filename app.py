import streamlit as st
import google.generativeai as genai
import dotenv
import os
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches

# Load environment variables
dotenv.load_dotenv()

# -------------------- Configuration --------------------
# Set API key for Google Gemini API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    st.error("Google API Key not found. Please set 'GOOGLE_API_KEY' in the .env file.")
    st.stop()

os.environ['GOOGLE_API_KEY'] = GOOGLE_API_KEY
genai.configure(api_key=GOOGLE_API_KEY)

# Model configuration settings
GENERATION_CONFIG = {
    'temperature': 1,
    'top_p': 0.95,
    'top_k': 0,
    'max_output_tokens': 8192
}

# Safety settings to filter out harmful content
SAFETY_SETTINGS = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
]

# System prompt for medical image analysis
SYSTEM_PROMPT = """
You are a domain expert in medical image analysis. You are tasked with 
examining medical images for a renowned hospital. Your expertise will help in 
identifying or discovering any anomalies, diseases, conditions, or health issues 
that might be present in the image.

Your key responsibilities:
1. Detailed Analysis: Scrutinize and thoroughly examine each image, focusing on finding abnormalities.
2. Analysis Report: Document findings and articulate them clearly in a structured format.
3. Recommendations: Suggest remedies, tests, or treatments based on the analysis.
4. Treatments: If applicable, lay out detailed treatments for faster recovery.

Important Notes:
1. Respond only if the image pertains to human health issues.
2. If the image is unclear, note: 'Unable to be correctly determined based on the uploaded image.'
3. Disclaimer: Accompany your analysis with: "Consult with a Doctor before making any decisions."
"""

# Initialize the generative model
model = genai.GenerativeModel(
    model_name="gemini-1.5-pro-latest",
    generation_config=GENERATION_CONFIG,
    safety_settings=SAFETY_SETTINGS
)

# -------------------- Streamlit UI Setup --------------------
# Configure Streamlit page
st.set_page_config(
    page_title='AI Medical Assistant',
    page_icon='./images/app.png',
    layout="wide"
)

# App Header
st.markdown("""
    <h1 style='text-align: center; font-size: 60px; line-height: 1.5; font-family: Arial, sans-serif;'>
        <span style='color: #008C74;'>AI</span> 
        <span style='color: #1E7F9B;'>Medical</span>
        <span style='color: #00C2A3;'>Assistant</span>
    </h1>
    <p style='text-align: center; font-size: 20px; color: #808080;'>
        Your virtual healthcare assistant for personalized health guidance.
    </p>
""", unsafe_allow_html=True)

# Sidebar for export button
with st.sidebar:
    st.header("Export and Options")
    export_button = st.button("Export as DOCX")
    clear_button = st.button("Clear All")

# File uploader for medical images
uploaded_file = st.file_uploader(
    'Please upload your medical image for analysis:', 
    type=['png', 'jpg', 'jpeg']
)

if uploaded_file:
    # Display uploaded image
    st.image(uploaded_file, width=200, caption='Uploaded Image')

# Generate report button
if st.button('Generate Report'):
    # Validate uploaded file
    if not uploaded_file:
        st.warning("Please upload an image before generating a report.")
    else:
        with st.spinner('Analyzing the image and generating your report...'):
            # Prepare image data for the model
            img_data = uploaded_file.getvalue()
            img_payload = {'mime_type': 'image/jpeg', 'data': img_data}

            # Prepare prompt input for the model
            prompt_input = [img_payload, SYSTEM_PROMPT]

            # Generate content using the AI model
            try:
                response = model.generate_content(prompt_input)
                if response:
                    st.success("Report Generated Successfully!")
                    report_text = response.text
                    st.write(report_text)

                    # Save report and image for export
                    if "report_content" not in st.session_state:
                        st.session_state.report_content = report_text
                    if "uploaded_image" not in st.session_state:
                        st.session_state.uploaded_image = uploaded_file

                else:
                    st.error("Failed to generate a report. Please try again.")
            except Exception as e:
                st.error(f"Error generating report: {str(e)}")

# Export as DOCX functionality
if export_button:
    if "report_content" in st.session_state and "uploaded_image" in st.session_state:
        # Create a Word document
        doc = Document()

        # Add heading
        doc.add_heading('AI Medical Assistant Report', level=1)

        # Add uploaded image
        img_path = Path("uploaded_image.jpg")
        with open(img_path, "wb") as img_file:
            img_file.write(st.session_state.uploaded_image.getvalue())
        doc.add_picture(str(img_path), width=Inches(4))

        # Add report content
        doc.add_paragraph(st.session_state.report_content)

        # Save the document
        doc_path = Path("AI_Medical_Report.docx")
        doc.save(doc_path)

        # Provide download link
        with open(doc_path, "rb") as file:
            st.download_button(
                label="Download Report as DOCX",
                data=file,
                file_name="AI_Medical_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Clean up temporary files
        img_path.unlink()
        doc_path.unlink()
    else:
        st.warning("No report available to export. Please generate a report first.")

# Clear session state and reset page
if clear_button:
    # Clear session state variables
    for key in list(st.session_state.keys()):
        del st.session_state[key]
    # Reset query params to reload the app
    st.query_params.clear()